# app.py

import os
import re
import concurrent.futures
import gradio as gr
from datetime import datetime, date
from docx import Document
from docx.shared import Pt, RGBColor, Inches

from agents.coordinator_agent import get_document_text
from agents.notes_agent import generate_notes
from agents.quiz_agent import generate_questions
from agents.planner_agent import generate_study_plan
from agents.youtube_agent import get_youtube_transcript
from agents.key_manager import validate_api_key, has_demo_key


# ══════════════════════════════════════════════════════════════════════════════
# DOCX HELPERS
# ══════════════════════════════════════════════════════════════════════════════

def _parse_inline(paragraph, text: str):
    pattern = r'(\*\*[^*\n]+?\*\*|\*[^*\n]+?\*)'
    for part in re.split(pattern, text):
        if not part:
            continue
        if part.startswith('**') and part.endswith('**') and len(part) > 4:
            paragraph.add_run(part[2:-2]).bold = True
        elif part.startswith('*') and part.endswith('*') and len(part) > 2:
            paragraph.add_run(part[1:-1]).italic = True
        else:
            paragraph.add_run(part)


def _is_table_row(line: str) -> bool:
    s = line.strip()
    return s.startswith('|') and s.endswith('|') and len(s) > 2


def _is_separator_row(line: str) -> bool:
    s = line.strip()
    return _is_table_row(s) and bool(re.match(r'^\|[\s\-:|]+\|$', s))


def _parse_table_cells(line: str):
    s = line.strip()
    if s.startswith('|'): s = s[1:]
    if s.endswith('|'):   s = s[:-1]
    return [cell.strip() for cell in s.split('|')]


def _add_table_to_doc(doc, rows):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    NAVY = RGBColor(0x1A, 0x17, 0x30)
    data_rows = [r for r in rows if not _is_separator_row(r)]
    if not data_rows:
        return
    col_count = max(len(_parse_table_cells(r)) for r in data_rows)
    table = doc.add_table(rows=len(data_rows), cols=col_count)
    table.style = 'Table Grid'
    for row_idx, raw_row in enumerate(data_rows):
        cells = _parse_table_cells(raw_row)
        while len(cells) < col_count:
            cells.append('')
        for col_idx, cell_text in enumerate(cells):
            cell = table.cell(row_idx, col_idx)
            para = cell.paragraphs[0]
            para.clear()
            is_header = (row_idx == 0)
            run = para.add_run(re.sub(r'\*+', '', cell_text).strip())
            run.bold = is_header
            run.font.size = Pt(10)
            run.font.color.rgb = NAVY
            if is_header:
                tc   = cell._tc
                tcPr = tc.get_or_add_tcPr()
                shd  = OxmlElement('w:shd')
                shd.set(qn('w:val'),   'clear')
                shd.set(qn('w:color'), 'auto')
                shd.set(qn('w:fill'),  'FEF3C7')
                tcPr.append(shd)


def save_docx(content: str, filename: str) -> str:
    doc = Document()
    for sec in doc.sections:
        sec.left_margin   = Inches(1.15)
        sec.right_margin  = Inches(1.15)
        sec.top_margin    = Inches(1.00)
        sec.bottom_margin = Inches(1.00)

    NAVY  = RGBColor(0x1A, 0x17, 0x30)
    NAVY2 = RGBColor(0x2D, 0x2A, 0x4A)
    NAVY3 = RGBColor(0x4A, 0x46, 0x6A)

    doc.styles['Heading 1'].font.size      = Pt(20)
    doc.styles['Heading 1'].font.bold      = True
    doc.styles['Heading 1'].font.color.rgb = NAVY
    doc.styles['Heading 2'].font.size      = Pt(15)
    doc.styles['Heading 2'].font.bold      = True
    doc.styles['Heading 2'].font.color.rgb = NAVY2
    doc.styles['Heading 3'].font.size      = Pt(12)
    doc.styles['Heading 3'].font.bold      = True
    doc.styles['Heading 3'].font.color.rgb = NAVY3
    doc.styles['Normal'].font.size         = Pt(11)
    doc.styles['Normal'].font.color.rgb    = NAVY

    lines = content.split('\n')
    i = 0
    while i < len(lines):
        line = lines[i]
        s    = line.strip()

        if _is_table_row(s):
            table_rows = []
            while i < len(lines) and _is_table_row(lines[i].strip()):
                table_rows.append(lines[i])
                i += 1
            _add_table_to_doc(doc, table_rows)
            doc.add_paragraph().paragraph_format.space_after = Pt(4)
            continue

        i += 1
        if not s:
            continue

        if s.startswith('#### '):
            doc.add_heading(re.sub(r'\*+', '', s[5:]).strip(), level=4)
        elif s.startswith('### '):
            doc.add_heading(re.sub(r'\*+', '', s[4:]).strip(), level=3)
        elif s.startswith('## '):
            doc.add_heading(re.sub(r'\*+', '', s[3:]).strip(), level=2)
        elif s.startswith('# '):
            doc.add_heading(re.sub(r'\*+', '', s[2:]).strip(), level=1)
        elif re.match(r'^[-*•]\s', s):
            _parse_inline(doc.add_paragraph(style='List Bullet'), s[2:].strip())
        elif re.match(r'^\d+[\.\)]\s', s):
            _parse_inline(doc.add_paragraph(style='List Number'),
                          re.sub(r'^\d+[\.\)]\s+', '', s))
        elif s in ('---', '***', '___'):
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(6)
        else:
            p = doc.add_paragraph()
            _parse_inline(p, s)
            p.paragraph_format.space_after = Pt(3)

    doc.save(filename)
    return filename


# ══════════════════════════════════════════════════════════════════════════════
# STAGE STEPPER
# ══════════════════════════════════════════════════════════════════════════════

STAGES = [
    ("📄", "Reading Input",            "Extracting text from your file or YouTube video"),
    ("⚡", "Notes + Plan in Parallel", "Notes Agent & Planner Agent running simultaneously"),
    ("❓", "Generating Questions",     "Question Agent building MCQ, essay & viva questions"),
    ("💾", "Saving Files",             "Writing properly formatted Word documents"),
]


def stage_html(active: int) -> str:
    items = []
    for i, (icon, title, desc) in enumerate(STAGES):
        if active < 0:
            state = "pending"
        elif i < active or active >= len(STAGES):
            state = "done"
        elif i == active:
            state = "active"
        else:
            state = "pending"
        dot_content = "✓" if state == "done" else icon
        items.append(f"""
            <div class="si si-{state}">
                <div class="si-dot">{dot_content}</div>
                <div class="si-body">
                    <b class="si-title">{title}</b>
                    <span class="si-desc">{desc}</span>
                </div>
            </div>""")
        if i < len(STAGES) - 1:
            cls = "cn-fill" if (active > i or active >= len(STAGES)) else "cn"
            items.append(f'<div class="{cls}"></div>')
    return f'<div class="sbar">{"".join(items)}</div>'


# ══════════════════════════════════════════════════════════════════════════════
# STATS
# ══════════════════════════════════════════════════════════════════════════════

def compute_stats(notes: str, questions: str, exam_date_str: str) -> str:
    topics = max(1, sum(1 for l in notes.split('\n') if l.strip().startswith('## ')))
    # FIX: question format changed from "1. ..." numbered lists to
    # "**Q:**" / "**Q1:**" style markers. The old regex r'^\d+[\.\)]' never
    # matched these, drastically undercounting (showed "6" when far more
    # were actually generated). Now matches any line starting with **Q
    # (covers **Q:, **Q1:, **Q2:, etc.)
    q_count = max(1, sum(1 for l in questions.split('\n') if re.match(r'^\*\*Q\d*:?\*\*', l.strip())))
    days = "—"
    if exam_date_str:
        try:
            exam_d  = datetime.strptime(exam_date_str, "%Y-%m-%d").date()
            today_d = date.today()
            delta   = (exam_d - today_d).days
            days    = max(0, delta)
        except Exception:
            pass
    return f"""
    <div class="stats-bar">
        <div class="stat-chip">
            <span class="stat-icon">📚</span>
            <span class="stat-val">{topics}</span>
            <span class="stat-lbl">Topics Covered</span>
        </div>
        <div class="stat-sep"></div>
        <div class="stat-chip">
            <span class="stat-icon">❓</span>
            <span class="stat-val">{q_count}</span>
            <span class="stat-lbl">Questions Generated</span>
        </div>
        <div class="stat-sep"></div>
        <div class="stat-chip">
            <span class="stat-icon">📅</span>
            <span class="stat-val">{days}</span>
            <span class="stat-lbl">Days Until Exam</span>
        </div>
    </div>"""


# ══════════════════════════════════════════════════════════════════════════════
# API KEY STATUS BADGE
# ══════════════════════════════════════════════════════════════════════════════

def key_status_html(mode: str, message: str = "") -> str:
    """
    mode: 'demo' | 'user' | 'invalid' | 'none'
    """
    configs = {
        "demo":    ("🟢", "Demo Mode", "Using the project's API key — free to test, shared quota.", "#D1FAE5", "#065F46", "rgba(16,185,129,0.25)"),
        "user":    ("🔑", "Your API Key", "Requests are now billed to your own Gemini account.", "#DBEAFE", "#1E40AF", "rgba(59,130,246,0.25)"),
        "invalid": ("🔴", "Invalid Key", message or "That key could not be validated.", "#FEE2E2", "#991B1B", "rgba(239,68,68,0.25)"),
        "none":    ("⚪", "No Key Available", "Paste a Gemini API key below to continue.", "#F3F4F6", "#374151", "rgba(107,114,128,0.25)"),
    }
    icon, title, desc, bg, fg, border = configs.get(mode, configs["none"])
    return f"""
    <div style="display:flex;align-items:center;gap:10px;background:{bg};border:1px solid {border};
                border-radius:10px;padding:10px 16px;margin-top:6px;">
        <span style="font-size:16px;">{icon}</span>
        <div style="display:flex;flex-direction:column;">
            <span style="font-size:12.5px;font-weight:700;color:{fg};">{title}</span>
            <span style="font-size:11px;color:{fg};opacity:0.85;">{desc}</span>
        </div>
    </div>"""


def on_key_change(user_key: str):
    """
    Called whenever the API key textbox changes.
    Validates the key live and updates the status badge.
    Returns (status_html, active_key_state)
    """
    user_key = (user_key or "").strip()

    if not user_key:
        # No user key entered — fall back to demo
        if has_demo_key():
            return key_status_html("demo"), ""
        return key_status_html("none"), ""

    is_valid, msg = validate_api_key(user_key)
    if is_valid:
        return key_status_html("user"), user_key
    else:
        # Invalid key entered — fall back to demo so the demo still works
        if has_demo_key():
            return key_status_html("invalid", msg + " Falling back to Demo Mode."), ""
        return key_status_html("invalid", msg), ""


# ══════════════════════════════════════════════════════════════════════════════
# MAIN PROCESSING FUNCTION
# ══════════════════════════════════════════════════════════════════════════════

_BLANK = ("", "", "", "", None, None, None)   # stats, notes, qs, plan, 3 files


def process_document(file, youtube_url, exam_date, hours_per_day, active_key):
    """
    active_key: the validated user API key (string) or "" to use the demo key.
    """
    executor = None
    try:
        has_file = file is not None
        has_url  = bool(youtube_url and youtube_url.strip())

        if not has_file and not has_url:
            yield (stage_html(-1),
                   "⚠️ Please upload a document OR paste a YouTube URL.",
                   *_BLANK)
            return

        if has_file and has_url:
            yield (stage_html(-1),
                   "⚠️ Please use either a file upload OR a YouTube URL — not both at once.",
                   *_BLANK)
            return

        # Resolve which key will actually be used for this run
        key_for_run = active_key.strip() if active_key and active_key.strip() else None
        if not key_for_run and not has_demo_key():
            yield (stage_html(-1),
                   "❌ Please enter a valid Gemini API key to continue.",
                   *_BLANK)
            return

        # Normalise exam date
        exam_date_str = ""
        if exam_date:
            try:
                if isinstance(exam_date, (int, float)):
                    exam_date_str = date.fromtimestamp(float(exam_date)).strftime("%Y-%m-%d")
                else:
                    raw = str(exam_date).split(" ")[0].split("T")[0]
                    datetime.strptime(raw, "%Y-%m-%d")
                    exam_date_str = raw
            except Exception:
                exam_date_str = ""

        # ── Stage 0: Extract text ──────────────────────────────────────────
        if has_url:
            yield (stage_html(0), "🎬 Fetching YouTube transcript…",
                   "", "", "", "", None, None, None)
            text = get_youtube_transcript(youtube_url.strip())
        else:
            yield (stage_html(0), "📄 Reading and parsing your document…",
                   "", "", "", "", None, None, None)
            text = get_document_text(file)

        if not text or not text.strip():
            yield (stage_html(-1),
                   "❌ Could not extract text. For PDFs check it is not a scanned image. "
                   "For YouTube check the video has captions enabled.",
                   *_BLANK)
            return

        # ── Stage 1: Notes + Plan in parallel ─────────────────────────────
        yield (stage_html(1),
               "⚡ Notes Agent & Planner Agent running simultaneously…",
               "", "", "", "", None, None, None)

        executor     = concurrent.futures.ThreadPoolExecutor(max_workers=2)
        notes_future = executor.submit(generate_notes, text, key_for_run)
        plan_future  = executor.submit(generate_study_plan, text, exam_date_str, hours_per_day, key_for_run)

        notes = notes_future.result()

        if not notes or not notes.strip():
            yield (stage_html(-1),
                   "❌ Notes Agent returned empty. Check your Gemini API key and quota.",
                   *_BLANK)
            executor.shutdown(wait=False)
            return

        # ── Stage 2: Questions ─────────────────────────────────────────────
        yield (stage_html(2),
               "✅ Notes complete! Question Agent now generating questions…",
               "", notes, "", "", None, None, None)

        qs_future  = executor.submit(generate_questions, notes, key_for_run)
        study_plan = plan_future.result()
        questions  = qs_future.result()
        executor.shutdown(wait=False)
        executor = None

        if not questions or not questions.strip():
            yield (stage_html(-1),
                   "❌ Quiz Agent returned empty. Check your Gemini API key and quota.",
                   "", notes, "", study_plan, None, None, None)
            return

        # ── Stage 3: Save DOCX ────────────────────────────────────────────
        yield (stage_html(3),
               "💾 Saving your formatted Word documents…",
               "", notes, questions, study_plan, None, None, None)

        os.makedirs("outputs", exist_ok=True)
        notes_f = save_docx(notes,      "outputs/notes.docx")
        qs_f    = save_docx(questions,  "outputs/questions.docx")
        plan_f  = save_docx(study_plan, "outputs/study_plan.docx")

        stats = compute_stats(notes, questions, exam_date_str)

        yield (
            stage_html(4),
            "✅ Your study materials are ready — explore the tabs below!",
            stats,
            notes, questions, study_plan,
            notes_f, qs_f, plan_f,
        )

    except Exception as e:
        if executor:
            executor.shutdown(wait=False)
        err = str(e)
        err_lower = err.lower()
        if "invalid_api_key" in err_lower:
            msg = "❌ Your Gemini API key is invalid. Please check it and try again."
        elif "quota" in err_lower or "resource_exhausted" in err_lower:
            msg = "❌ Gemini API quota exceeded — please try again later or use your own API key."
        elif "no gemini api key" in err_lower:
            msg = "❌ Please enter a valid Gemini API key to continue."
        elif "transcript" in err_lower or "caption" in err_lower or "disabled" in err_lower:
            msg = f"❌ YouTube error: {err}"
        else:
            msg = f"❌ Error: {err}"
        yield (stage_html(-1), msg, *_BLANK)


# ══════════════════════════════════════════════════════════════════════════════
# CSS
# ══════════════════════════════════════════════════════════════════════════════

CSS = """
@import url('https://fonts.googleapis.com/css2?family=Playfair+Display:wght@600;700&family=Inter:wght@300;400;500;600&display=swap');

:root {
    --gold:        #C8890A;
    --gold-hi:     #E8A020;
    --gold-pastel: #FEF3C7;
    --gold-border: rgba(200,137,10,0.22);
    --gold-glow:   rgba(200,137,10,0.18);
    --page:        #FDFAF3;
    --surface:     #FFFFFF;
    --surface-2:   #FBF8F0;
    --border:      #EAE4D2;
    --border-hi:   rgba(200,137,10,0.40);
    --ink:         #1A1730;
    --ink-soft:    #3C3860;
    --muted:       #7A7490;
    --yt-red:      #FF0000;
    --yt-dark:     #CC0000;
    --radius:      14px;
    --shadow:      0 2px 12px rgba(0,0,0,0.07);
}

*, *::before, *::after { box-sizing: border-box; }

body, .gradio-container {
    background: var(--page) !important;
    font-family: 'Inter', system-ui, sans-serif !important;
    color: var(--ink) !important;
}

.progress-bar-wrap, .progress-level, .progress-level-inner,
.eta-bar, [data-testid="progress-bar"] { display: none !important; }

.block, .gr-group, .gap {
    background: var(--surface) !important;
    border: 1px solid var(--border) !important;
    border-radius: var(--radius) !important;
    box-shadow: var(--shadow) !important;
}

/* ─── API key box ─── */
.key-box textarea, .key-box input {
    border: 2px solid #C7D2FE !important;
    background: #F5F6FF !important;
    border-radius: 10px !important;
    font-family: monospace !important;
}
.key-box textarea:focus, .key-box input:focus {
    border-color: #6366F1 !important;
    box-shadow: 0 0 0 3px rgba(99,102,241,0.12) !important;
}

/* ─── OR divider ─── */
.or-divider {
    display: flex; align-items: center; gap: 12px;
    color: var(--muted); font-size: 12px; font-weight: 600;
    letter-spacing: 0.08em; margin: 8px 0;
}
.or-divider::before, .or-divider::after {
    content: ''; flex: 1; height: 1px; background: var(--border);
}

/* ─── YouTube input ─── */
.yt-box textarea, .yt-box input {
    border: 2px solid #FECACA !important;
    background: #FFF5F5 !important;
    border-radius: 10px !important;
}
.yt-box textarea:focus, .yt-box input:focus {
    border-color: var(--yt-red) !important;
    box-shadow: 0 0 0 3px rgba(255,0,0,0.10) !important;
}

/* ─── Stage Stepper ─── */
.sbar {
    display: flex; align-items: flex-start; gap: 0; padding: 18px 20px;
    background: var(--surface) !important;
    border: 1px solid var(--border) !important;
    border-top: 3px solid var(--gold-hi) !important;
    border-radius: var(--radius) !important;
    box-shadow: var(--shadow) !important;
    margin: 6px 0;
}
.si { display: flex; align-items: flex-start; gap: 10px; flex: 1; min-width: 0; }
.si-dot {
    width: 32px; height: 32px; border-radius: 50%;
    display: flex; align-items: center; justify-content: center;
    font-size: 15px; flex-shrink: 0;
    border: 2px solid var(--border);
    background: var(--surface-2); color: var(--muted);
    transition: all 0.3s ease;
}
.si-body  { display: flex; flex-direction: column; gap: 3px; min-width: 0; }
.si-title { font-size: 11.5px; font-weight: 600; letter-spacing: 0.03em; color: var(--muted); }
.si-desc  { font-size: 10.5px; color: var(--muted); opacity: 0.65; line-height: 1.3; }
.cn, .cn-fill {
    height: 2px; width: 28px; margin-top: 18px;
    flex-shrink: 0; border-radius: 2px; transition: background 0.4s ease;
}
.cn      { background: var(--border); }
.cn-fill { background: linear-gradient(90deg, var(--gold), var(--gold-hi)); }
.si-pending .si-dot   { filter: opacity(0.45); }
.si-pending .si-title { opacity: 0.4; }
.si-active .si-dot {
    background: var(--gold-pastel); border-color: var(--gold-hi); color: var(--gold);
    animation: si-pulse 1s ease-in-out infinite alternate; font-size: 18px;
}
.si-active .si-title { color: var(--gold) !important; }
.si-done .si-dot {
    background: var(--gold-hi); border-color: var(--gold);
    color: white; font-size: 14px; font-weight: 700;
}
.si-done .si-title { color: var(--gold) !important; }
@keyframes si-pulse {
    from { transform: scale(1);    box-shadow: 0 0 0   0   rgba(200,137,10,0); }
    to   { transform: scale(1.06); box-shadow: 0 0 0 6px rgba(200,137,10,0.18); }
}

/* ─── Status message ─── */
#status-md .prose p, #status-md p {
    font-size: 13.5px !important; font-weight: 600 !important;
    color: var(--ink) !important; background: var(--gold-pastel) !important;
    border-left: 4px solid var(--gold-hi) !important;
    border-radius: 8px !important; padding: 11px 16px !important; margin: 0 !important;
}

/* ─── Stats bar ─── */
.stats-bar {
    display: flex; align-items: center; justify-content: center;
    background: var(--surface); border: 1px solid var(--gold-border);
    border-radius: var(--radius); padding: 14px 20px; box-shadow: var(--shadow);
}
.stat-chip { display: flex; align-items: center; gap: 8px; flex: 1; justify-content: center; }
.stat-icon { font-size: 20px; }
.stat-val  { font-size: 22px; font-weight: 700; color: var(--gold); font-family: 'Playfair Display', serif; }
.stat-lbl  { font-size: 12px; color: var(--muted); font-weight: 500; }
.stat-sep  { width: 1px; height: 36px; background: var(--border); flex-shrink: 0; margin: 0 12px; }

/* ─── Output panels ─── */
.out-panel {
    background: var(--surface) !important;
    border: 1px solid var(--border) !important;
    border-radius: var(--radius) !important;
    padding: 6px 8px !important;
    box-shadow: var(--shadow) !important;
}
.out-panel .prose {
    --tw-prose-body: var(--ink); --tw-prose-headings: #0F0D1E;
    --tw-prose-bold: var(--ink); --tw-prose-bullets: var(--gold-hi);
    --tw-prose-counters: var(--gold-hi); --tw-prose-links: var(--gold);
    --tw-prose-invert-body: var(--ink); --tw-prose-invert-headings: #0F0D1E;
    --tw-prose-invert-bold: var(--ink); --tw-prose-invert-bullets: var(--gold-hi);
    background: transparent !important; max-width: none !important; padding: 16px 20px !important;
}
.out-panel .prose h1, .out-panel .prose h2 {
    color: #0F0D1E !important; font-family: 'Playfair Display', serif !important;
    border-bottom: 2px solid var(--gold-hi) !important; padding-bottom: 4px !important;
}
.out-panel .prose h3, .out-panel .prose h4 {
    color: #2B284A !important; font-family: 'Playfair Display', serif !important;
}
.out-panel .prose p, .out-panel .prose li, .out-panel .prose td,
.out-panel .prose th, .out-panel .prose strong { color: var(--ink) !important; }
.out-panel .prose strong { font-weight: 700 !important; }
.out-panel .prose code {
    background: #F0E9D5 !important; color: var(--ink) !important;
    border-radius: 4px !important; padding: 1px 6px !important; font-size: 0.88em !important;
}
.out-panel .prose blockquote {
    border-left: 3px solid var(--gold-hi) !important; color: var(--ink-soft) !important;
    font-style: italic; background: var(--gold-pastel) !important;
    border-radius: 0 6px 6px 0; padding: 8px 14px !important;
}

/* ─── Generate button ─── */
#gen-btn {
    background: linear-gradient(135deg, #E8A020 0%, #B87010 100%) !important;
    border: 1px solid rgba(200,137,10,0.45) !important;
    color: #FFFFFF !important; font-weight: 700 !important; font-size: 15px !important;
    letter-spacing: 0.05em !important; border-radius: 10px !important;
    padding: 14px 28px !important; width: 100% !important;
    box-shadow: 0 4px 20px var(--gold-glow) !important;
    transition: transform 0.15s, box-shadow 0.15s !important;
}
#gen-btn:hover  { transform: translateY(-2px) !important; box-shadow: 0 8px 32px var(--gold-glow) !important; }
#gen-btn:active { transform: translateY(0) !important; }

/* ─── Tabs ─── */
.tab-nav button {
    color: var(--muted) !important; background: transparent !important;
    border: none !important; border-bottom: 2px solid transparent !important;
    font-weight: 500 !important; font-size: 13.5px !important;
    padding: 10px 18px !important; border-radius: 0 !important; transition: color 0.2s !important;
}
.tab-nav button.selected {
    color: var(--gold) !important; border-bottom: 2px solid var(--gold-hi) !important; font-weight: 600 !important;
}
.tab-nav button:hover:not(.selected) { color: var(--ink) !important; }
#tab-notes { border-top: 3px solid #E8A020 !important; }
#tab-qs    { border-top: 3px solid #8B5CF6 !important; }
#tab-plan  { border-top: 3px solid #10B981 !important; }

/* ─── Upload zone ─── */
.upload-zone .wrap {
    border: 2px dashed var(--border-hi) !important;
    background: var(--gold-pastel) !important;
    border-radius: var(--radius) !important; transition: background 0.2s !important;
}
.upload-zone .wrap:hover { background: #FDE68A !important; }

/* ─── Inputs ─── */
label span, .label-wrap span {
    color: var(--ink-soft) !important; font-size: 13px !important; font-weight: 500 !important;
}
input[type="number"], input[type="text"], input[type="date"], textarea, input[type="password"] {
    background: var(--surface-2) !important; border: 1px solid var(--border) !important;
    color: var(--ink) !important; border-radius: 8px !important;
}
input:focus, textarea:focus {
    border-color: var(--gold-hi) !important;
    box-shadow: 0 0 0 3px var(--gold-glow) !important; outline: none !important;
}

/* ─── Section labels ─── */
.section-label {
    color: var(--muted); font-size: 11px; font-weight: 600;
    letter-spacing: 0.10em; text-transform: uppercase;
    margin: 20px 0 8px; display: flex; align-items: center; gap: 8px;
}
.section-label::after { content: ''; flex: 1; height: 1px; background: var(--border); }

/* ─── Scrollbars ─── */
::-webkit-scrollbar       { width: 5px; height: 5px; }
::-webkit-scrollbar-track { background: var(--surface-2); }
::-webkit-scrollbar-thumb { background: var(--gold-hi); border-radius: 3px; }
"""

# ══════════════════════════════════════════════════════════════════════════════
# THEME
# ══════════════════════════════════════════════════════════════════════════════

THEME = gr.themes.Base(
    primary_hue="orange", neutral_hue="zinc",
    font=[gr.themes.GoogleFont("Inter"), "ui-sans-serif", "system-ui"],
).set(
    body_background_fill="#FDFAF3",        body_background_fill_dark="#FDFAF3",
    body_text_color="#1A1730",             body_text_color_dark="#1A1730",
    block_background_fill="#FFFFFF",       block_background_fill_dark="#FFFFFF",
    block_border_color="#EAE4D2",          block_border_color_dark="#EAE4D2",
    block_title_text_color="#1A1730",      block_title_text_color_dark="#1A1730",
    input_background_fill="#FBF8F0",       input_background_fill_dark="#FBF8F0",
    input_border_color="#EAE4D2",          input_border_color_dark="#EAE4D2",
    input_placeholder_color="#9CA3AF",
    button_primary_background_fill="linear-gradient(135deg,#E8A020,#B87010)",
    button_primary_text_color="#FFFFFF",
)

# ══════════════════════════════════════════════════════════════════════════════
# UI
# ══════════════════════════════════════════════════════════════════════════════

with gr.Blocks(theme=THEME, css=CSS) as demo:

    # State holding the currently ACTIVE (validated) user key, or "" for demo
    active_key = gr.State(value="")

    # ── Hero ──────────────────────────────────────────────────────────────────
    gr.HTML("""
    <div style="padding:32px 0 20px;text-align:center;">
      <div style="display:inline-block;background:linear-gradient(160deg,#FFFBF0 0%,#FEF9EC 60%,#FFF8E8 100%);border:1px solid rgba(200,137,10,0.30);border-top:4px solid #E8A020;border-radius:20px;padding:28px 56px 24px;box-shadow:0 4px 32px rgba(200,137,10,0.12);">
          <div style="font-size:52px;margin-bottom:8px;line-height:1;">🎓</div>
          <h1 style="margin:0;font-family:'Playfair Display',serif;font-size:2.3rem;font-weight:700;color:#1A1730;letter-spacing:-0.02em;line-height:1.1;">StudyMate <span style="color:#C8890A;">AI</span></h1>
          <p style="margin:8px 0 0;color:#7A7490;font-size:13px;letter-spacing:0.08em;text-transform:uppercase;font-weight:400;">AI-Powered Study Companion</p>
          <div style="display:flex;gap:10px;justify-content:center;margin-top:16px;flex-wrap:wrap;">
              <span style="background:#FEF3C7;border:1px solid rgba(200,137,10,0.30);color:#92640A;padding:5px 14px;border-radius:20px;font-size:12px;font-weight:600;">📚 Smart Notes</span>
              <span style="background:#EDE9FE;border:1px solid rgba(124,58,237,0.20);color:#5B21B6;padding:5px 14px;border-radius:20px;font-size:12px;font-weight:600;">❓ Practice Questions</span>
              <span style="background:#D1FAE5;border:1px solid rgba(16,185,129,0.22);color:#065F46;padding:5px 14px;border-radius:20px;font-size:12px;font-weight:600;">📅 Study Planner</span>
              <span style="background:#FEE2E2;border:1px solid rgba(255,0,0,0.20);color:#CC0000;padding:5px 14px;border-radius:20px;font-size:12px;font-weight:600;">🎬 YouTube Notes</span>
          </div>
      </div>
    </div>
    """)

    # ── API Key Section ──────────────────────────────────────────────────────
    gr.HTML('<div class="section-label">🔑 API Key Settings</div>')
    with gr.Row():
        with gr.Column(scale=3):
            default_key = os.environ.get("GEMINI_API_KEY", "")
            api_key_input = gr.Textbox(
                label="Your Gemini API Key (optional)",
                placeholder="Paste your own Gemini API key here, or leave blank to use Demo Mode",
                type="password",
                elem_classes=["key-box"],
            )
            gr.HTML(
                '<p style="color:#7A7490;font-size:11.5px;margin:4px 0 0;line-height:1.4;">'
                'Don\'t have a key? Get one free at '
                '<a href="https://aistudio.google.com/apikey" target="_blank" style="color:#6366F1;">'
                'aistudio.google.com/apikey</a>. Leaving this blank uses the project\'s shared demo key.'
                '</p>'
            )
        with gr.Column(scale=2):
            key_status_out = gr.HTML(
                value=key_status_html("demo" if has_demo_key() else "none")
            )

    # ── Architecture ──────────────────────────────────────────────────────────
    gr.HTML("""
    <div style="margin:16px 0 16px;">
        <p style="color:#7A7490;font-size:11px;text-align:center;letter-spacing:0.10em;text-transform:uppercase;margin-bottom:12px;font-weight:600;">Multi-Agent Architecture</p>
        <div style="display:flex;gap:8px;align-items:stretch;">
            <div style="flex:1;background:#FFFBF0;border:1px solid #EAE4D2;border-top:3px solid #E8A020;border-radius:14px;padding:16px 12px;text-align:center;box-shadow:0 2px 8px rgba(0,0,0,0.05);">
                <div style="font-size:26px;margin-bottom:6px;">📄🎬</div>
                <div style="color:#92640A;font-weight:700;font-size:11px;letter-spacing:0.06em;margin-bottom:4px;">COORDINATOR</div>
                <div style="color:#7A7490;font-size:11px;line-height:1.4;">File or YouTube transcript</div>
            </div>
            <div style="display:flex;align-items:center;color:#EAE4D2;font-size:18px;flex-shrink:0;">→</div>
            <div style="flex:2;background:#FFFBF0;border:1px solid #EAE4D2;border-top:3px solid #E8A020;border-radius:14px;padding:14px 12px;box-shadow:0 2px 8px rgba(0,0,0,0.05);">
                <div style="text-align:center;color:#92640A;font-weight:700;font-size:11px;letter-spacing:0.06em;margin-bottom:10px;">⚡ PARALLEL</div>
                <div style="display:flex;gap:8px;">
                    <div style="flex:1;background:#FEF3C7;border:1px solid rgba(200,137,10,0.25);border-radius:10px;padding:10px 8px;text-align:center;">
                        <div style="font-size:20px;">📚</div>
                        <div style="color:#92640A;font-weight:600;font-size:10.5px;margin-top:4px;">Notes Agent</div>
                        <div style="color:#7A7490;font-size:10px;margin-top:2px;">Summarises content</div>
                    </div>
                    <div style="flex:1;background:#D1FAE5;border:1px solid rgba(16,185,129,0.22);border-radius:10px;padding:10px 8px;text-align:center;">
                        <div style="font-size:20px;">📅</div>
                        <div style="color:#065F46;font-weight:600;font-size:10.5px;margin-top:4px;">Planner Agent</div>
                        <div style="color:#7A7490;font-size:10px;margin-top:2px;">Builds schedule</div>
                    </div>
                </div>
            </div>
            <div style="display:flex;align-items:center;color:#EAE4D2;font-size:18px;flex-shrink:0;">→</div>
            <div style="flex:1;background:#FFFBF0;border:1px solid rgba(124,58,237,0.18);border-top:3px solid #8B5CF6;border-radius:14px;padding:16px 12px;text-align:center;box-shadow:0 2px 8px rgba(0,0,0,0.05);">
                <div style="font-size:26px;margin-bottom:6px;">❓</div>
                <div style="color:#5B21B6;font-weight:700;font-size:11px;letter-spacing:0.06em;margin-bottom:4px;">QUESTIONS</div>
                <div style="color:#7A7490;font-size:10px;margin-top:4px;font-style:italic;">(serial after notes)</div>
            </div>
        </div>
        <p style="color:#A09AB8;font-size:10.5px;text-align:center;margin-top:10px;font-style:italic;">
            Notes + Plan run simultaneously → Questions start as soon as Notes finish
        </p>
    </div>
    """)

    # ── Inputs ────────────────────────────────────────────────────────────────
    gr.HTML('<div class="section-label">Step 1 — Upload Document or Paste YouTube Link</div>')
    with gr.Row(equal_height=False):
        with gr.Column(scale=3):
            file_input = gr.File(
                label="📎 Upload document (PDF, PPTX or DOCX)",
                type="filepath",
                file_types=[".pdf", ".pptx", ".docx"],
                elem_classes=["upload-zone"],
            )
            gr.HTML('<div class="or-divider">OR</div>')
            youtube_url = gr.Textbox(
                label="🎬 Paste YouTube Lecture URL",
                placeholder="https://www.youtube.com/watch?v=...",
                info="Video must have captions/subtitles enabled",
                elem_classes=["yt-box"],
            )
        with gr.Column(scale=2):
            exam_date = gr.DateTime(
                label="📅 Exam / Assessment Date",
                info="When do you need to be exam-ready?",
                value=datetime.today().strftime("%Y-%m-%d"),
                include_time=False,
            )
            hours = gr.Number(
                label="⏱ Study Hours Available Per Day",
                value=3, minimum=0.5, maximum=16, step=0.5,
            )

    # ── Progress ──────────────────────────────────────────────────────────────
    gr.HTML('<div class="section-label">Step 2 — Live Generation Progress</div>')
    stage_out  = gr.HTML(value=stage_html(-1))
    status_out = gr.Markdown(value="", elem_id="status-md")

    gen_btn = gr.Button("⚡ Generate My Study Materials", variant="primary", elem_id="gen-btn")

    stats_out = gr.HTML(value="")

    # ── Output tabs ───────────────────────────────────────────────────────────
    gr.HTML('<div class="section-label">Step 3 — Your Study Materials</div>')
    with gr.Tabs():
        with gr.Tab("📚  Study Notes", elem_id="tab-notes"):
            gr.HTML('<p style="color:#7A7490;font-size:12.5px;margin:4px 0 10px;line-height:1.5;">AI-condensed notes — structured with headings, bullets, and key concepts.</p>')
            notes_out  = gr.Markdown(elem_classes=["out-panel"])
            notes_file = gr.File(label="⬇ Download Study Notes (.docx)", type="filepath")

        with gr.Tab("❓  Practice Questions", elem_id="tab-qs"):
            gr.HTML('<p style="color:#7A7490;font-size:12.5px;margin:4px 0 10px;line-height:1.5;">MCQs, essay prompts, and viva questions generated from your notes.</p>')
            qs_out  = gr.Markdown(elem_classes=["out-panel"])
            qs_file = gr.File(label="⬇ Download Questions (.docx)", type="filepath")

        with gr.Tab("📅  Study Plan", elem_id="tab-plan"):
            gr.HTML('<p style="color:#7A7490;font-size:12.5px;margin:4px 0 10px;line-height:1.5;">A personalised day-by-day schedule across your available study time.</p>')
            plan_out  = gr.Markdown(elem_classes=["out-panel"])
            plan_file = gr.File(label="⬇ Download Study Plan (.docx)", type="filepath")

    # ── Event binding: API key validation (live, on change) ────────────────────
    api_key_input.change(
        fn=on_key_change,
        inputs=[api_key_input],
        outputs=[key_status_out, active_key],
    )

    # ── Event binding: Generate ─────────────────────────────────────────────
    gen_btn.click(
        fn=process_document,
        show_progress="hidden",
        inputs=[file_input, youtube_url, exam_date, hours, active_key],
        outputs=[
            stage_out, status_out, stats_out,
            notes_out, qs_out, plan_out,
            notes_file, qs_file, plan_file,
        ],
    )

    gr.HTML("""
    <div style="text-align:center;padding:22px 0 10px;margin-top:16px;border-top:1px solid #EAE4D2;">
        <span style="color:#9CA3AF;font-size:12px;letter-spacing:0.04em;">
            💛 StudyMate AI &nbsp;·&nbsp; Multi-Agent Learning Assistant &nbsp;·&nbsp; Powered by Gemini
        </span>
    </div>
    """)

demo.launch(share=True)

