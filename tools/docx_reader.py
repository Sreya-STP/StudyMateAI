from docx import Document

def extract_docx_text(docx_path):
    doc = Document(docx_path)
    text = ""

    # 1. Extract standard paragraphs
    for para in doc.paragraphs:
        if para.text.strip():
            text += para.text + "\n"

    # 2. Extract text from tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    text += cell.text + " | " # Add a pipe to separate cell data
            text += "\n" # Newline at the end of each row

    return text